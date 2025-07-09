import azure.functions as func
import logging
import traceback
import inspect
import psycopg2
from psycopg2 import sql
from datetime import datetime
from db_config import DB_HOST, DB_NAME, DB_USER, DB_PASSWORD, DB_PORT
from azure.storage.blob import BlobServiceClient
from azure.identity import DefaultAzureCredential
import openai
import os
import pandas as pd
import json
from docx import Document
import mimetypes
import uuid
import time
from openai import AzureOpenAI
from typing import Dict, List, Any, Optional, Tuple, Union
import xml.etree.ElementTree as ET
import io
from functools import lru_cache
from contextlib import contextmanager

app = func.FunctionApp(http_auth_level=func.AuthLevel.ANONYMOUS)

# Enhanced logging helper functions with improved line number tracking and traceback
def get_caller_info(frame_depth: int = 2):
    """Get detailed caller information including filename, function name, and line number"""
    try:
        frame = inspect.currentframe()
        # Navigate to the appropriate frame depth
        for _ in range(frame_depth):
            frame = frame.f_back
        
        filename = os.path.basename(frame.f_code.co_filename)
        function_name = frame.f_code.co_name
        line_number = frame.f_lineno
        
        return f"{filename}:{function_name}:L{line_number}"
    except Exception:
        return "unknown:unknown:L0"

def log_info(message: str):
    """Log info message with detailed caller information"""
    caller_info = get_caller_info()
    logging.info(f"[{caller_info}] {message}")

def log_warning(message: str, include_traceback: bool = False):
    """Log warning message with caller information and optional traceback"""
    caller_info = get_caller_info()
    if include_traceback:
        tb_str = traceback.format_exc()
        logging.warning(f"[{caller_info}] {message}\nTraceback:\n{tb_str}")
    else:
        logging.warning(f"[{caller_info}] {message}")

def log_error(message: str, include_traceback: bool = True):
    """Log error message with caller information and traceback (enabled by default)"""
    caller_info = get_caller_info()
    if include_traceback:
        tb_str = traceback.format_exc()
        logging.error(f"[{caller_info}] {message}\nTraceback:\n{tb_str}")
    else:
        logging.error(f"[{caller_info}] {message}")

def log_exception(message: str, exception: Exception = None):
    """Log exception with full traceback and caller information"""
    caller_info = get_caller_info()
    if exception:
        tb_str = ''.join(traceback.format_exception(type(exception), exception, exception.__traceback__))
        logging.error(f"[{caller_info}] {message}\nException: {type(exception).__name__}: {str(exception)}\nFull Traceback:\n{tb_str}")
    else:
        tb_str = traceback.format_exc()
        logging.error(f"[{caller_info}] {message}\nCurrent Traceback:\n{tb_str}")

def log_debug(message: str):
    """Log debug message with detailed caller information"""
    caller_info = get_caller_info()
    logging.debug(f"[{caller_info}] {message}")

def log_critical(message: str, include_traceback: bool = True):
    """Log critical message with caller information and traceback"""
    caller_info = get_caller_info()
    if include_traceback:
        tb_str = traceback.format_exc()
        logging.critical(f"[{caller_info}] {message}\nTraceback:\n{tb_str}")
    else:
        logging.critical(f"[{caller_info}] {message}")

# Utility functions for safe JSON operations
def safe_json_loads(data: Union[str, dict, list], default: Any = None) -> Any:
    """Safely parse JSON data with fallback"""
    if data is None:
        return default if default is not None else {}
    
    if isinstance(data, (dict, list)):
        return data
    
    if isinstance(data, str):
        try:
            return json.loads(data)
        except json.JSONDecodeError as e:
            log_exception(f"JSON decode error for data: {data[:100]}...", e)
            return default if default is not None else {}
    
    return default if default is not None else {}

def safe_json_dumps(data: Any) -> str:
    """Safely convert data to JSON string"""
    try:
        return json.dumps(data) if data is not None else "{}"
    except (TypeError, ValueError) as e:
        log_exception(f"JSON encode error", e)
        return "{}"

@contextmanager
def get_database_connection():
    """Context manager for database connections with proper cleanup"""
    conn = None
    try:
        conn = psycopg2.connect(
            host=DB_HOST,
            dbname=DB_NAME,
            user=DB_USER,
            password=DB_PASSWORD,
            port=DB_PORT,
            sslmode='require'
        )
        yield conn
    except Exception as e:
        log_exception(f"Database connection failed", e)
        if conn:
            conn.rollback()
        raise
    finally:
        if conn:
            conn.close()

class OrderParsingService:
    """Enhanced service for parsing orders using Azure OpenAI"""
    
    def __init__(self):
        self.openai_endpoint = os.environ.get("AZURE_OPENAI_ENDPOINT")
        self.openai_deployment = os.environ.get("AZURE_OPENAI_DEPLOYMENT", "gpt-4o")
        self.openai_key = os.environ.get("AZURE_OPENAI_KEY")
        self.openai_version = os.environ.get("AZURE_OPENAI_VERSION", "2024-11-20")
        self.client = None
        self._initialize_client()
    
    def _initialize_client(self):
        """Initialize Azure OpenAI client with error handling"""
        if not self.openai_endpoint:
            log_warning("Azure OpenAI endpoint not configured. AI-powered parsing will be disabled.")
            return
        
        try:
            self.client = AzureOpenAI(
                azure_endpoint=self.openai_endpoint,
                api_key=self.openai_key,
                api_version=self.openai_version,
            )
            self._test_connection()
            log_info("Azure OpenAI client initialized successfully")
        except Exception as e:
            log_exception(f"Failed to initialize Azure OpenAI client", e)
            self.client = None
    
    def _test_connection(self):
        """Test Azure OpenAI connection"""
        try:
            response = self.client.chat.completions.create(
                model=self.openai_deployment,
                messages=[{"role": "user", "content": "test"}],
                max_tokens=5,
                temperature=0
            )
            log_info("Azure OpenAI connection test successful")
        except Exception as e:
            log_exception(f"Azure OpenAI connection test failed", e)
            raise
    
    def _make_api_call(self, messages: List[Dict], max_tokens: int = 1500, temperature: float = 0.1) -> Optional[str]:
        """Make API call with error handling and retries"""
        if not self.client:
            return None
        
        max_retries = 3
        for attempt in range(max_retries):
            try:
                response = self.client.chat.completions.create(
                    model=self.openai_deployment,
                    messages=messages,
                    temperature=temperature,
                    max_tokens=max_tokens,
                    response_format={"type": "json_object"}
                )
                return response.choices[0].message.content
            
            except openai.APIConnectionError as e:
                log_exception(f"Azure OpenAI connection error (attempt {attempt + 1})", e)
                if attempt < max_retries - 1:
                    time.sleep(2 ** attempt)  # Exponential backoff
                continue
            except openai.RateLimitError as e:
                log_exception(f"Azure OpenAI rate limit error (attempt {attempt + 1})", e)
                if attempt < max_retries - 1:
                    time.sleep(5)
                continue
            except openai.APIStatusError as e:
                log_exception(f"Azure OpenAI API error: {e.status_code} - {e.response}", e)
                break
            except Exception as e:
                log_exception(f"Unexpected error in API call", e)
                break
        
        return None
    
    def analyze_order_completeness(self, parsed_data: Dict[str, Any], file_type: str) -> Dict[str, Any]:
        """Analyze order data for completeness using Azure OpenAI"""
        if not self.client:
            return self._get_fallback_analysis("AI analysis unavailable")
        
        try:
            prompt = self._create_analysis_prompt(parsed_data, file_type)
            messages = [
                {"role": "system", "content": "You are an expert in analyzing order data for FMCG supply chain operations. Provide detailed analysis in JSON format."},
                {"role": "user", "content": prompt}
            ]
            
            response_content = self._make_api_call(messages, max_tokens=1500)
            if response_content:
                return safe_json_loads(response_content, self._get_fallback_analysis("Invalid response"))
            else:
                return self._get_fallback_analysis("API call failed")
                
        except Exception as e:
            log_exception(f"Azure OpenAI analysis failed", e)
            return self._get_fallback_analysis(f"Analysis error: {str(e)}")
    
    def extract_sku_items(self, parsed_data: Dict[str, Any], file_type: str) -> List[Dict[str, Any]]:
        """Extract SKU items from parsed data using AI"""
        if not self.client:
            return []
        
        try:
            prompt = self._create_sku_extraction_prompt(parsed_data, file_type)
            messages = [
                {"role": "system", "content": "You are an expert in extracting product/SKU information from order data. Return a JSON array of SKU items."},
                {"role": "user", "content": prompt}
            ]
            
            response_content = self._make_api_call(messages, max_tokens=2000)
            if not response_content:
                return []
            
            response_data = safe_json_loads(response_content, {})
            sku_items = response_data.get("sku_items", []) if isinstance(response_data, dict) else response_data
            
            # Validate SKU items
            if isinstance(sku_items, list):
                validated_items = []
                for item in sku_items:
                    if isinstance(item, dict):
                        # Ensure required fields have default values
                        validated_item = {
                            "sku_code": item.get("sku_code", ""),
                            "product_name": item.get("product_name", ""),
                            "category": item.get("category"),
                            "brand": item.get("brand", ""),
                            "quantity_ordered": max(0, int(item.get("quantity_ordered", 0))),
                            "unit_of_measure": item.get("unit_of_measure"),
                            "unit_price": item.get("unit_price"),
                            "total_price": item.get("total_price"),
                            "weight_kg": item.get("weight_kg"),
                            "volume_m3": item.get("volume_m3"),
                            "temperature_requirement": item.get("temperature_requirement"),
                            "fragile": bool(item.get("fragile", False)),
                            "product_attributes": item.get("product_attributes", {}),
                            "processing_remarks": item.get("processing_remarks", "")
                        }
                        validated_items.append(validated_item)
                    else:
                        log_warning(f"SKU item is not a dictionary, skipping: {type(item)}")
                return validated_items
            else:
                log_warning(f"SKU items is not a list: {type(sku_items)}")
                return []
                
        except Exception as e:
            log_exception(f"SKU extraction failed", e)
            return []
    
    def extract_retailer_information(self, parsed_data: Dict[str, Any], file_type: str) -> Dict[str, Any]:
        """Extract retailer information from order data using Azure OpenAI"""
        if not self.client:
            return {
                "retailer_extracted": False,
                "reason": "AI analysis unavailable",
                "extracted_info": {}
            }
        
        try:
            prompt = self._create_retailer_extraction_prompt(parsed_data, file_type)
            messages = [
                {"role": "system", "content": "You are an expert in extracting retailer information from order documents. Extract retailer details and return in JSON format."},
                {"role": "user", "content": prompt}
            ]
            
            response_content = self._make_api_call(messages, max_tokens=1000)
            if response_content:
                return safe_json_loads(response_content, {
                    "retailer_extracted": False,
                    "reason": "Invalid response format",
                    "extracted_info": {}
                })
            else:
                return {
                    "retailer_extracted": False,
                    "reason": "API call failed",
                    "extracted_info": {}
                }
                
        except Exception as e:
            log_exception(f"Error extracting retailer information", e)
            return {
                "retailer_extracted": False,
                "reason": f"Extraction error: {str(e)}",
                "extracted_info": {}
            }

    def _get_fallback_analysis(self, error_reason: str) -> Dict[str, Any]:
        """Return fallback analysis when AI analysis fails"""
        return {
            "completeness_score": 0.5,
            "missing_fields": [f"AI analysis unavailable: {error_reason}"],
            "validation_errors": [],
            "recommendations": ["Manual review required due to AI analysis failure"],
            "order_summary": {
                "total_sku_count": 0,
                "estimated_total_quantity": 0,
                "has_pricing": False,
                "has_delivery_info": False
            }
        }
    
    def _create_analysis_prompt(self, parsed_data: Dict[str, Any], file_type: str) -> str:
        """Create prompt for order completeness analysis"""
        # Truncate large data to avoid token limits
        data_str = safe_json_dumps(parsed_data)
        if len(data_str) > 10000:
            data_str = data_str[:10000] + "... [truncated]"
        
        return f"""
        Analyze the following {file_type} order data for completeness and quality:

        Data: {data_str}

        Required fields for a complete order:
        - Order identification (order number, reference)
        - Customer/retailer information
        - Delivery details (address, requested delivery date)
        - Product items with SKU codes, quantities, and descriptions
        - Pricing information (unit prices, totals)
        - Special instructions or requirements

        Please provide analysis in this JSON format:
        {{
            "completeness_score": <float between 0.0 and 1.0>,
            "missing_fields": ["list of missing required fields"],
            "validation_errors": ["list of data quality issues"],
            "recommendations": ["list of improvement suggestions"],
            "order_summary": {{
                "total_sku_count": <number>,
                "estimated_total_quantity": <number>,
                "has_pricing": <boolean>,
                "has_delivery_info": <boolean>
            }}
        }}
        """
    
    def _create_sku_extraction_prompt(self, parsed_data: Dict[str, Any], file_type: str) -> str:
        """Create prompt for SKU item extraction"""
        
        # Handle text-based formats specially
        if file_type in ["Text", "Log", "XML", "Word Document"]:
            content_field_map = {
                "Text": "full_content",
                "Log": "full_content", 
                "XML": "full_content",
                "Word Document": "full_text"
            }
            
            content_field = content_field_map.get(file_type)
            if content_field and content_field in parsed_data:
                text_content = parsed_data.get(content_field, "")
                # Truncate very large content
                if isinstance(text_content, str) and len(text_content) > 15000:
                    text_content = text_content[:15000] + "... [content truncated]"
                
                return f"""
                You are an expert in extracting product order information from unstructured {file_type} data.
                
                Below is the content that contains order information:
                
                {file_type} Content:
                ```
                {text_content}
                ```
                
                Extract ALL product items from this content. Look for:
                - Product codes/SKUs
                - Product names/descriptions  
                - Quantities
                - Prices
                - Units/measurements
                - Any other product attributes
                
                Return your findings as a JSON object with a "sku_items" array:
                {{
                    "sku_items": [
                        {{
                            "sku_code": "product SKU or identifier",
                            "product_name": "product description/name", 
                            "category": "product category if available",
                            "brand": "brand name if available",
                            "quantity_ordered": <integer>,
                            "unit_of_measure": "unit type",
                            "unit_price": <decimal or null>,
                            "total_price": <decimal or null>,
                            "weight_kg": <decimal or null>,
                            "volume_m3": <decimal or null>,
                            "temperature_requirement": "ambient/chilled/frozen or null",
                            "fragile": <boolean>,
                            "product_attributes": {{"any": "additional attributes"}},
                            "processing_remarks": "List any missing details or confidence notes"
                        }}
                    ]
                }}
                
                Guidelines:
                - Use null for unknown fields instead of making up values
                - Default quantity_ordered to 1 if not specified but item clearly appears ordered
                - Include processing_remarks explaining confidence level and missing info
                - Return empty array if no order items found
                """
        
        # Default handling for structured formats
        data_str = safe_json_dumps(parsed_data)
        if len(data_str) > 10000:
            data_str = data_str[:10000] + "... [truncated]"
            
        return f"""
        Extract individual SKU/product items from this {file_type} order data:

        Data: {data_str}

        Return a JSON object with a "sku_items" array where each item has this structure:
        {{
            "sku_items": [
                {{
                    "sku_code": "product SKU or identifier",
                    "product_name": "product description/name",
                    "category": "product category if available", 
                    "brand": "brand name if available",
                    "quantity_ordered": <integer>,
                    "unit_of_measure": "unit type",
                    "unit_price": <decimal or null>,
                    "total_price": <decimal or null>,
                    "weight_kg": <decimal or null>,
                    "volume_m3": <decimal or null>,
                    "temperature_requirement": "ambient/chilled/frozen or null",
                    "fragile": <boolean>,
                    "product_attributes": {{"any": "additional attributes"}},
                    "processing_remarks": "List any missing required details"
                }}
            ]
        }}

        Guidelines:
        - Use null for unavailable data, note missing info in processing_remarks
        - Ensure quantity_ordered is always a positive integer
        - Always return valid JSON with the "sku_items" array
        - For processing_remarks, note missing critical info like SKU code, pricing, category, etc.
        """
    
    def _create_retailer_extraction_prompt(self, parsed_data: Dict[str, Any], file_type: str) -> str:
        """Create prompt for retailer information extraction"""
        data_str = safe_json_dumps(parsed_data)
        if len(data_str) > 10000:
            data_str = data_str[:10000] + "... [truncated]"
            
        return f"""
        Extract retailer information from this {file_type} order data:

        Data: {data_str}

        Return a JSON object with retailer information:
        {{
            "retailer_extracted": <boolean>,
            "confidence_score": <float 0-1>,
            "extracted_info": {{
                "retailer_name": "company/store name",
                "retailer_code": "retailer identifier/code if available",
                "contact_person": "contact person name if available",
                "contact_email": "email address if available",
                "contact_phone": "phone number if available",
                "delivery_address": {{
                    "street": "street address",
                    "city": "city name", 
                    "state": "state/province",
                    "postal_code": "zip/postal code",
                    "country": "country"
                }},
                "business_details": {{
                    "tax_id": "tax identification if available",
                    "business_type": "retail type if mentioned",
                    "store_number": "store identifier if available"
                }}
            }},
            "extraction_notes": "notes about what was found or missing"
        }}

        Look for retailer information in:
        - Company names, store names, business names
        - "Bill to", "Ship to", "Delivery to" addresses
        - Contact information (emails, phones)
        - Business identifiers (codes, tax IDs, store numbers)

        Set retailer_extracted to true only if you find a clear retailer name and at least one additional piece of information.
        """

def parse_file_content(file_data: bytes, file_extension: str, filename: str) -> Dict[str, Any]:
    """Enhanced file parsing with better error handling"""
    try:
        if file_extension in ['.xlsx', '.xls']:
            return parse_excel_file(file_data)
        elif file_extension in ['.csv']:
            return parse_csv_file(file_data)
        elif file_extension in ['.json']:
            return parse_json_file(file_data)
        elif file_extension in ['.xml']:
            return parse_xml_file(file_data)
        elif file_extension in ['.txt', '.log']:
            return parse_text_file(file_data, file_extension)
        elif file_extension in ['.docx']:
            return parse_docx_file(file_data)
        else:
            return {
                "file_type": "Unsupported",
                "file_extension": file_extension,
                "file_size": len(file_data)
            }
    except Exception as e:
        log_exception(f"Error parsing file content", e)
        return {
            "file_type": "Error",
            "error": str(e),
            "file_extension": file_extension
        }

def parse_excel_file(file_data: bytes) -> Dict[str, Any]:
    """Parse Excel file using pandas"""
    try:
        excel_file = pd.ExcelFile(io.BytesIO(file_data))
        sheets_data = {}
        
        for sheet_name in excel_file.sheet_names[:5]:  # Limit to first 5 sheets
            df = pd.read_excel(excel_file, sheet_name=sheet_name)
            sheets_data[sheet_name] = {
                "columns": df.columns.tolist(),
                "row_count": len(df),
                "sample_data": df.head(10).to_dict('records'),
                "data": df.to_dict('records') if len(df) <= 1000 else df.head(1000).to_dict('records')
            }
        
        return {
            "file_type": "Excel",
            "sheet_names": excel_file.sheet_names,
            "sheet_count": len(excel_file.sheet_names),
            "sheets_data": sheets_data
        }
    except Exception as e:
        log_error(f"Excel parsing error: {e}")
        raise

def parse_csv_file(file_data: bytes) -> Dict[str, Any]:
    """Parse CSV file using pandas"""
    try:
        df = pd.read_csv(io.BytesIO(file_data))
        return {
            "file_type": "CSV",
            "columns": df.columns.tolist(),
            "row_count": len(df),
            "sample_data": df.head(10).to_dict('records'),
            "data": df.to_dict('records') if len(df) <= 1000 else df.head(1000).to_dict('records')
        }
    except Exception as e:
        log_error(f"CSV parsing error: {e}")
        raise

def parse_json_file(file_data: bytes) -> Dict[str, Any]:
    """Parse JSON file"""
    try:
        json_data = json.loads(file_data.decode('utf-8'))
        if isinstance(json_data, list):
            return {
                "file_type": "JSON",
                "structure": "array",
                "item_count": len(json_data),
                "sample_data": json_data[:10] if len(json_data) > 0 else [],
                "keys": list(json_data[0].keys()) if len(json_data) > 0 and isinstance(json_data[0], dict) else [],
                "all_data": json_data
            }
        else:
            return {
                "file_type": "JSON",
                "structure": "object", 
                "keys": list(json_data.keys()) if isinstance(json_data, dict) else [],
                "sample_data": json_data,
                "all_data": json_data
            }
    except Exception as e:
        log_error(f"JSON parsing error: {e}")
        raise

def parse_xml_file(file_data: bytes) -> Dict[str, Any]:
    """Parse XML file"""
    try:
        root = ET.fromstring(file_data.decode('utf-8', errors='replace'))
        
        def traverse_element(element, path=""):
            current_path = path + "/" + element.tag if path else element.tag
            element_data = {"tag": element.tag, "attributes": element.attrib}
            
            if element.text and element.text.strip():
                element_data["text"] = element.text.strip()
                
            children = list(element)
            if children:
                element_data["children"] = [traverse_element(child, current_path) for child in children]
            
            return element_data
        
        xml_structure = traverse_element(root)
        
        return {
            "file_type": "XML",
            "root_tag": root.tag,
            "structure": xml_structure,
            "full_content": file_data.decode('utf-8', errors='replace'),
            "llm_processing_required": True
        }
    except ET.ParseError as xml_error:
        log_error(f"XML parsing error: {xml_error}")
        return {
            "file_type": "XML",
            "error": f"Invalid XML format: {str(xml_error)}",
            "full_content": file_data.decode('utf-8', errors='replace'),
            "llm_processing_required": True
        }

def parse_text_file(file_data: bytes, file_extension: str) -> Dict[str, Any]:
    """Parse text or log file"""
    try:
        text_content = file_data.decode('utf-8', errors='replace')
        lines = text_content.split('\n')
        
        result = {
            "file_type": "Text" if file_extension == '.txt' else "Log",
            "line_count": len(lines),
            "character_count": len(text_content),
            "word_count": len(text_content.split()),
            "sample_lines": lines[:20],
            "full_content": text_content,
            "llm_processing_required": True
        }
        
        # Detect structure hints
        structure_hints = []
        order_keywords = ["order", "sku", "product", "quantity", "delivery", "item", "price"]
        
        for keyword in order_keywords:
            if keyword.lower() in text_content.lower():
                structure_hints.append(f"Contains '{keyword}' references")
        
        # Pattern detection
        if any(line.strip().startswith("{") and line.strip().endswith("}") for line in lines):
            structure_hints.append("Contains potential JSON-like entries")
        if any(line.count(":") > 1 for line in lines):
            structure_hints.append("Contains colon-separated entries")
        if any(line.count(",") > 2 for line in lines):
            structure_hints.append("Contains comma-separated entries")
        if any(line.count("|") > 2 for line in lines):
            structure_hints.append("Contains pipe-separated entries")
            
        if structure_hints:
            result["structure_hints"] = structure_hints
            
        return result
    except Exception as e:
        log_error(f"Text parsing error: {e}")
        raise

def parse_docx_file(file_data: bytes) -> Dict[str, Any]:
    """Parse Word document"""
    try:
        doc = Document(io.BytesIO(file_data))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables_data = []
        
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables_data.append(table_data)
        
        return {
            "file_type": "Word Document",
            "paragraph_count": len(paragraphs),
            "sample_paragraphs": paragraphs[:10],
            "tables_count": len(doc.tables),
            "tables_data": tables_data,
            "full_text": '\n'.join(paragraphs),
            "llm_processing_required": True
        }
    except Exception as e:
        log_error(f"DOCX parsing error: {e}")
        raise

def insert_order_tracking(conn, order_id: str, status: str, message: str, details: Dict[str, Any]):
    """Insert order tracking entry with enhanced error handling"""
    try:
        cur = conn.cursor()
        insert_query = sql.SQL("""
            INSERT INTO order_tracking (id, order_id, status, message, details, created_at)
            VALUES (%s, %s, %s, %s, %s, NOW())
        """)
        cur.execute(insert_query, (
            str(uuid.uuid4()),
            order_id,
            status,
            message,
            safe_json_dumps(details)
        ))
        cur.close()
    except Exception as e:
        log_error(f"Failed to insert order tracking: {e}")
        raise

def insert_sku_items(conn, order_id: str, sku_items: List[Dict[str, Any]]):
    """Insert SKU items with better error handling and validation"""
    if not sku_items:
        log_info(f"No SKU items to insert for order {order_id}")
        return
    
    try:
        cur = conn.cursor()
        
        # Delete existing SKU items
        delete_query = sql.SQL("DELETE FROM order_sku_items WHERE order_id = %s")
        cur.execute(delete_query, (order_id,))
        
        # Insert new SKU items
        insert_query = sql.SQL("""
            INSERT INTO order_sku_items (
                id, order_id, sku_code, product_name, category, brand,
                quantity_ordered, unit_of_measure, unit_price, total_price,
                weight_kg, volume_m3, temperature_requirement, fragile,
                product_attributes, processing_remarks, created_at, updated_at
            ) VALUES (
                %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, NOW(), NOW()
            )
        """)
        
        successful_inserts = 0
        for item in sku_items:
            try:
                if not isinstance(item, dict):
                    log_warning(f"SKU item is not a dictionary: {type(item)}")
                    continue
                
                # Validate and clean data
                quantity_ordered = max(0, int(item.get('quantity_ordered', 0))) if item.get('quantity_ordered') is not None else 0
                unit_price = float(item.get('unit_price')) if item.get('unit_price') is not None else None
                total_price = float(item.get('total_price')) if item.get('total_price') is not None else None
                weight_kg = float(item.get('weight_kg')) if item.get('weight_kg') is not None else None
                volume_m3 = float(item.get('volume_m3')) if item.get('volume_m3') is not None else None
                
                cur.execute(insert_query, (
                    str(uuid.uuid4()),
                    order_id,
                    str(item.get('sku_code', ''))[:255],  # Truncate if too long
                    str(item.get('product_name', ''))[:500],  # Truncate if too long
                    item.get('category'),
                    str(item.get('brand', ''))[:255],
                    quantity_ordered,
                    item.get('unit_of_measure'),
                    unit_price,
                    total_price,
                    weight_kg,
                    volume_m3,
                    item.get('temperature_requirement'),
                    bool(item.get('fragile', False)),
                    safe_json_dumps(item.get('product_attributes', {})),
                    str(item.get('processing_remarks', ''))[:1000],  # Truncate if too long
                ))
                successful_inserts += 1
                
            except (ValueError, TypeError) as data_error:
                log_warning(f"Data validation error for SKU item {item}: {data_error}")
                continue
            except Exception as item_error:
                log_error(f"Failed to insert individual SKU item {item}: {item_error}")
                continue
        
        cur.close()
        log_info(f"Successfully inserted {successful_inserts}/{len(sku_items)} SKU items for order {order_id}")
        
    except Exception as e:
        log_error(f"Failed to insert SKU items: {e}")
        raise

def update_order_summary(conn, order_id: str, analysis_result: Dict[str, Any], sku_items: List[Dict[str, Any]]):
    """Update order with analysis results and summary"""
    try:
        cur = conn.cursor()
        
        # Calculate summary statistics safely
        total_sku_count = len(sku_items)
        total_quantity = sum(max(0, int(item.get('quantity_ordered', 0))) for item in sku_items)
        total_weight = sum(float(item.get('weight_kg', 0) or 0) for item in sku_items)
        total_volume = sum(float(item.get('volume_m3', 0) or 0) for item in sku_items)
        
        # Calculate subtotal safely
        subtotal = 0
        for item in sku_items:
            try:
                price = float(item.get('total_price', 0) or 0)
                subtotal += price
            except (ValueError, TypeError):
                continue
        
        # Determine new order status
        completeness_score = analysis_result.get('completeness_score', 0.0)
        missing_fields = analysis_result.get('missing_fields', [])
        validation_errors = analysis_result.get('validation_errors', [])
        
        if total_sku_count > 0:
            if completeness_score >= 0.9:
                new_status = "READY_FOR_ASSIGNMENT"
            elif completeness_score >= 0.7:
                new_status = "NEEDS_REVIEW"
            elif completeness_score >= 0.5:
                new_status = "PENDING_REVIEW"
            else:
                new_status = "INCOMPLETE"
        else:
            new_status = "UPLOADED"
        
        log_info(f"Updating order {order_id} status to {new_status} (completeness: {completeness_score:.2f})")
        
        update_query = sql.SQL("""
            UPDATE orders SET
                parsed_data = %s,
                missing_fields = %s,
                validation_errors = %s,
                total_sku_count = %s,
                total_quantity = %s,
                total_weight_kg = %s,
                total_volume_m3 = %s,
                subtotal = %s,
                status = %s,
                updated_at = NOW()
            WHERE id = %s
        """)
        
        cur.execute(update_query, (
            safe_json_dumps(analysis_result.get('order_summary', {})),
            safe_json_dumps(missing_fields),
            safe_json_dumps(validation_errors),
            total_sku_count,
            total_quantity,
            total_weight,
            total_volume,
            subtotal,
            new_status,
            order_id
        ))
        
        cur.close()
        
    except Exception as e:
        log_error(f"Failed to update order summary: {e}")
        raise

@app.route(route="order_file_reader")
def order_file_reader(req: func.HttpRequest) -> func.HttpResponse:
    """Enhanced order file reader with AI-powered parsing and completeness checking"""
    log_info('Enhanced order file reader function processed a request.')
    
    # Extract order_id from request
    order_id = req.params.get('order_id')
    if not order_id:
        try:
            req_body = req.get_json()
            if req_body:
                order_id = req_body.get('order_id')
        except ValueError:
            pass

    if not order_id:
        return func.HttpResponse(
            "Please provide order_id in the query string or in the request body.",
            status_code=400
        )

    parsing_service = OrderParsingService()
    
    try:
        with get_database_connection() as conn:
            cur = conn.cursor()
            
            # Query order details
            query = sql.SQL("""
                SELECT id, order_number, priority, requested_delivery_date, file_path, 
                       status, original_filename, file_type, file_size
                FROM orders
                WHERE id = %s
            """)
            cur.execute(query, (order_id,))
            order_details = cur.fetchone()

            if not order_details or not order_details[4]:  # file_path doesn't exist
                return func.HttpResponse(
                    "No file_path found for this order or order not found.",
                    status_code=404
                )

            order_uuid, order_number, priority, delivery_date, file_path, status, filename, file_type, file_size = order_details
            
            # Process the file
            parse_status, parse_message, parsed_data, analysis_result, sku_items = process_order_file(
                file_path, parsing_service
            )
            
            # Insert results into database
            details = create_tracking_details(parsed_data, analysis_result, sku_items, file_path)
            insert_order_tracking(conn, order_id, parse_status, parse_message, details)
            
            if parse_status not in ["PARSING_FAILED"]:
                insert_sku_items(conn, order_id, sku_items)
                update_order_summary(conn, order_id, analysis_result, sku_items)
            
            conn.commit()
            cur.close()
        
        # Prepare response
        result = create_response_data(
            order_id, order_details, parse_status, parse_message, 
            parsed_data, analysis_result, sku_items
        )
        
        return func.HttpResponse(
            safe_json_dumps(result),
            status_code=200,
            mimetype="application/json"
        )
        
    except Exception as e:
        log_error(f"Database error: {e}", include_traceback=True)
        return func.HttpResponse(
            f"Database error: {e}",
            status_code=500
        )

def process_order_file(file_path: str, parsing_service: OrderParsingService) -> Tuple[str, str, Dict, Dict, List]:
    """Process order file and return results"""
    try:
        blob_connection_str = os.environ.get("REQUESTED_ORDERS_BLOB_CONNECTION_STRING")
        
        if "/" not in file_path:
            raise ValueError("file_path must be in the format 'container/blobname'")
        
        logging.info(f"Processing file at path: {file_path} , connection string: {blob_connection_str}")

        container = "requestedorders"
        _, blob_name = file_path.split(f"{container}/", 1)
        
        blob_service_client = BlobServiceClient.from_connection_string(blob_connection_str)
        blob_client = blob_service_client.get_blob_client(container=container, blob=blob_name)
        
        log_info(f"Processing file: {blob_name} in container: {container}")
        
        # Download and parse file
        file_data = blob_client.download_blob().readall()
        file_extension = os.path.splitext(blob_name)[1].lower()
        parsed_data = parse_file_content(file_data, file_extension, blob_name)
        
        if parsed_data.get("file_type") == "Unsupported":
            return "PARSING_FAILED", f"Unsupported file type: {file_extension}", {}, {}, []
        
        if parsed_data.get("file_type") == "Error":
            return "PARSING_FAILED", f"File parsing error: {parsed_data.get('error')}", {}, {}, []
        
        # AI analysis
        analysis_result = parsing_service.analyze_order_completeness(
            parsed_data, parsed_data.get("file_type", "Unknown")
        )
        
        # Extract SKU items
        sku_items = parsing_service.extract_sku_items(
            parsed_data, parsed_data.get("file_type", "Unknown")
        )
        
        # Determine status and message
        completeness_score = analysis_result.get('completeness_score', 0.0)
        if completeness_score >= 0.8:
            parse_status = "PARSING_COMPLETE"
            parse_message = f"Order parsed successfully with {completeness_score:.1%} completeness"
        elif completeness_score >= 0.5:
            parse_status = "PARSING_PARTIAL"
            parse_message = f"Order parsed with {completeness_score:.1%} completeness. Some fields may be missing."
        else:
            parse_status = "PARSING_INCOMPLETE"
            parse_message = f"Order parsing incomplete ({completeness_score:.1%}). Manual review required."
        
        log_info(f"Extracted {len(sku_items)} SKU items from {file_extension} file")
        
        return parse_status, parse_message, parsed_data, analysis_result, sku_items
        
    except Exception as e:
        log_error(f"Failed to process file: {e}", include_traceback=True)
        return "PARSING_FAILED", f"Failed to parse file: {e}", {"error": str(e)}, {}, []

def create_tracking_details(parsed_data: Dict, analysis_result: Dict, sku_items: List, file_path: str) -> Dict:
    """Create tracking details dictionary"""
    return {
        "parsing_results": parsed_data,
        "ai_analysis": analysis_result,
        "sku_items_count": len(sku_items),
        "file_path": file_path
    }

def create_response_data(order_id: str, order_details: tuple, parse_status: str, 
                        parse_message: str, parsed_data: Dict, analysis_result: Dict, 
                        sku_items: List) -> Dict:
    """Create response data dictionary"""
    order_uuid, order_number, priority, delivery_date, file_path, status, filename, file_type, file_size = order_details
    
    return {
        "order_id": order_id,
        "order_number": order_number,
        "priority": priority,
        "requested_delivery_date": delivery_date.strftime('%Y-%m-%d') if isinstance(delivery_date, datetime) else str(delivery_date) if delivery_date else None,
        "file_info": {
            "original_filename": filename,
            "file_path": file_path,
            "file_type": file_type,
            "file_size": file_size
        },
        "parse_status": parse_status,
        "parse_message": parse_message,
        "parsed_data": parsed_data,
        "ai_analysis": analysis_result,
        "sku_items": sku_items,
        "summary": {
            "total_sku_count": len(sku_items),
            "completeness_score": analysis_result.get('completeness_score', 0.0),
            "missing_fields_count": len(analysis_result.get('missing_fields', [])),
            "validation_errors_count": len(analysis_result.get('validation_errors', []))
        }
    }

@app.route(route="draft_order_email")
def draft_order_email(req: func.HttpRequest) -> func.HttpResponse:
    """Draft email for retailer issues or FMCG notification - FIXED"""
    log_info('Order email drafting function processed a request.')
    
    order_id = req.params.get('order_id')
    if not order_id:
        try:
            req_body = req.get_json()
            if req_body:
                order_id = req_body.get('order_id')
        except ValueError:
            pass

    if not order_id:
        return func.HttpResponse(
            "Please provide order_id in the query string or in the request body.",
            status_code=400
        )

    try:
        with get_database_connection() as conn:
            cur = conn.cursor()
            
            # Get order details
            query = sql.SQL("""
                SELECT o.id, o.order_number, o.parsed_data, o.missing_fields, 
                       o.validation_errors, o.total_sku_count, o.status,
                       COUNT(osi.id) as actual_sku_count,
                       o.retailer_id, r.name as retailer_name, r.contact_email
                FROM orders o
                LEFT JOIN order_sku_items osi ON o.id = osi.order_id
                LEFT JOIN retailers r ON o.retailer_id = r.id
                WHERE o.id = %s
                GROUP BY o.id, o.order_number, o.parsed_data, o.missing_fields, 
                         o.validation_errors, o.total_sku_count, o.status,
                         o.retailer_id, r.name, r.contact_email
            """)
            cur.execute(query, (order_id,))
            row = cur.fetchone()
            
            if not row:
                return func.HttpResponse(
                    f"Order with id {order_id} not found.",
                    status_code=404
                )
            
            # Extract and safely parse order details
            order_uuid, order_number, parsed_data, missing_fields, validation_errors, \
            total_sku_count, status, actual_sku_count, retailer_id, retailer_name, contact_email = row
            
            # FIXED: Safe parsing of JSON fields
            missing_fields_list = safe_json_loads(missing_fields, [])
            validation_errors_list = safe_json_loads(validation_errors, [])
            parsed_data_dict = safe_json_loads(parsed_data, {})
            
            # Determine email type and content
            has_issues = len(missing_fields_list) > 0 or len(validation_errors_list) > 0
            
            if has_issues:
                # Draft retailer email about issues
                email_type = "RETAILER_ISSUES"
                recipient = contact_email or "retailer@example.com"
                subject = f"Important: Issues Found with Order {order_number}"
                body = draft_retailer_issues_email(
                    order_number=order_number,
                    missing_fields=missing_fields_list,
                    validation_errors=validation_errors_list,
                    retailer_name=retailer_name
                )
            else:
                # Draft FMCG notification email
                email_type = "FMCG_NOTIFICATION"
                recipient = os.environ.get("FMCG_NOTIFICATION_EMAIL", "fmcg@example.com")
                subject = f"New Complete Order {order_number} Ready for Processing"
                body = draft_fmcg_notification_email(
                    order_number=order_number,
                    retailer_name=retailer_name,
                    sku_count=actual_sku_count,
                    status=status,
                    parsed_data=parsed_data_dict
                )
            
            # Create email record
            email_id = str(uuid.uuid4())
            sender = os.environ.get("SYSTEM_EMAIL_SENDER", "orders@orderplanner.com")
            
            insert_query = sql.SQL("""
                INSERT INTO email_communications (
                    id, order_id, email_type, recipient, sender, subject, body, 
                    created_at, status
                ) VALUES (
                    %s, %s, %s, %s, %s, %s, %s, NOW(), %s
                )
            """)
            
            cur.execute(insert_query, (
                email_id, order_id, email_type, recipient, sender, subject, body, "pending"
            ))
            
            conn.commit()
            cur.close()
        
        return func.HttpResponse(
            safe_json_dumps({
                "email_id": email_id,
                "email_type": email_type,
                "recipient": recipient,
                "subject": subject,
                "status": "pending",
                "created_at": datetime.now().isoformat()
            }),
            status_code=200,
            mimetype="application/json"
        )
        
    except Exception as e:
        log_error(f"Email drafting error: {e}", include_traceback=True)
        return func.HttpResponse(
            f"Email drafting error: {e}",
            status_code=500
        )

def draft_retailer_issues_email(order_number: str, missing_fields: List[str], 
                               validation_errors: List[str], retailer_name: str = None) -> str:
    """FIXED: Draft an email for retailers when their order has issues"""
    current_date = datetime.now().strftime("%B %d, %Y")
    greeting = f"Dear {retailer_name}" if retailer_name else "Dear Valued Retailer"
    
    # FIXED: Ensure inputs are lists
    if not isinstance(missing_fields, list):
        missing_fields = [str(missing_fields)] if missing_fields else []
    
    if not isinstance(validation_errors, list):
        validation_errors = [str(validation_errors)] if validation_errors else []
    
    # Build missing fields section
    missing_fields_html = ""
    if missing_fields and any(field for field in missing_fields if field):
        missing_fields_html = "<h3>Missing Information:</h3><ul>"
        for field in missing_fields:
            if field and field.strip():  # Only add non-empty fields
                missing_fields_html += f"<li>{str(field).strip()}</li>"
        missing_fields_html += "</ul>"
    
    # Build validation errors section
    validation_errors_html = ""
    if validation_errors and any(error for error in validation_errors if error):
        validation_errors_html = "<h3>Data Quality Issues:</h3><ul>"
        for error in validation_errors:
            if error and str(error).strip():  # Only add non-empty errors
                validation_errors_html += f"<li>{str(error).strip()}</li>"
        validation_errors_html += "</ul>"
    
    email_body = f"""
    <html>
    <head>
        <style>
            body {{ font-family: Arial, sans-serif; line-height: 1.6; }}
            .container {{ max-width: 600px; margin: 0 auto; padding: 20px; }}
            .header {{ background-color: #f8f9fa; padding: 15px; border-bottom: 3px solid #0078d4; }}
            .content {{ padding: 20px 0; }}
            .footer {{ font-size: 12px; color: #666; margin-top: 30px; padding-top: 10px; border-top: 1px solid #eee; }}
            .btn {{ display: inline-block; background-color: #0078d4; color: white; padding: 10px 20px; text-decoration: none; border-radius: 4px; }}
            h3 {{ color: #333; }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header">
                <h2>Order Review Required</h2>
            </div>
            <div class="content">
                <p>{greeting},</p>
                
                <p>We have received your order <strong>#{order_number}</strong> and found some issues that need your attention before we can process it.</p>
                
                {missing_fields_html}
                {validation_errors_html}
                
                <p>Please log in to your account to review and update this order, or reply to this email with the required information.</p>
                
                <p><a href="#" class="btn">Review Order</a></p>
                
                <p>If you have any questions or need assistance, please contact our support team.</p>
                
                <p>Thank you for your business.</p>
                
                <p>Sincerely,<br>Order Processing Team</p>
            </div>
            <div class="footer">
                <p>This is an automated message. Please do not reply directly to this email.</p>
                <p>Date: {current_date}</p>
            </div>
        </div>
    </body>
    </html>
    """
    
    return email_body

def draft_fmcg_notification_email(order_number: str, retailer_name: str = None, 
                                 sku_count: int = 0, status: str = None, 
                                 parsed_data: Dict[str, Any] = None) -> str:
    """FIXED: Draft a notification email for FMCG team about a complete order"""
    current_date = datetime.now().strftime("%B %d, %Y")
    retailer_info = retailer_name if retailer_name else "Unknown Retailer"
    
    # FIXED: Safe type conversions
    try:
        sku_count = int(sku_count) if sku_count is not None else 0
    except (ValueError, TypeError):
        sku_count = 0
    
    status = str(status) if status else "Unknown"
    
    if not isinstance(parsed_data, dict):
        parsed_data = {}
    
    # Build order summary
    order_summary = f"""
    <h3>Order Summary:</h3>
    <ul>
        <li><strong>Order Number:</strong> {order_number}</li>
        <li><strong>Retailer:</strong> {retailer_info}</li>
        <li><strong>SKU Count:</strong> {sku_count}</li>
        <li><strong>Status:</strong> {status}</li>
    """
    
    if parsed_data.get('has_delivery_info'):
        order_summary += "<li><strong>Delivery Information:</strong> Available</li>"
    
    order_summary += "</ul>"
    
    email_body = f"""
    <html>
    <head>
        <style>
            body {{ font-family: Arial, sans-serif; line-height: 1.6; }}
            .container {{ max-width: 600px; margin: 0 auto; padding: 20px; }}
            .header {{ background-color: #f0f7ff; padding: 15px; border-bottom: 3px solid #0078d4; }}
            .content {{ padding: 20px 0; }}
            .footer {{ font-size: 12px; color: #666; margin-top: 30px; padding-top: 10px; border-top: 1px solid #eee; }}
            .btn {{ display: inline-block; background-color: #0078d4; color: white; padding: 10px 20px; text-decoration: none; border-radius: 4px; }}
            h3 {{ color: #333; }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header">
                <h2>[INTERNAL] New Order Ready for Processing</h2>
            </div>
            <div class="content">
                <p>Hello FMCG Team,</p>
                
                <p>A new order has been validated and is ready for processing.</p>
                
                {order_summary}
                
                <p>This order has passed all validation checks and can be processed according to standard procedures.</p>
                
                <p><a href="#" class="btn">View Order Details</a></p>
                
                <p>Please process this order according to our standard operating procedures.</p>
                
                <p>Thank you,<br>Order Management System</p>
            </div>
            <div class="footer">
                <p>This is an automated internal notification. Please do not share outside the organization.</p>
                <p>Date: {current_date}</p>
            </div>
        </div>
    </body>
    </html>
    """
    
    return email_body

@app.route(route="health")
def health_check(req: func.HttpRequest) -> func.HttpResponse:
    """Health check endpoint"""
    return func.HttpResponse(
        safe_json_dumps({
            "status": "healthy",
            "timestamp": datetime.now().isoformat(),
            "azure_openai_configured": bool(os.environ.get("AZURE_OPENAI_ENDPOINT"))
        }),
        status_code=200,
        mimetype="application/json"
    )