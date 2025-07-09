import azure.functions as func
import logging
import psycopg2
from psycopg2 import sql
from datetime import datetime
from db_config import DB_HOST, DB_NAME, DB_USER, DB_PASSWORD, DB_PORT
from azure.storage.blob import BlobServiceClient
from azure.identity import DefaultAzureCredential, get_bearer_token_provider
import openai
import os
import pandas as pd
import json
from docx import Document
import mimetypes
import uuid
import time
from openai import AzureOpenAI
from typing import Dict, List, Any, Optional, Tuple

app = func.FunctionApp(http_auth_level=func.AuthLevel.ANONYMOUS)

class OrderParsingService:
    """Service for parsing orders using Azure OpenAI with recommended authentication pattern"""
    
    def __init__(self):
        # Initialize Azure OpenAI client using Microsoft's recommended pattern
        self.openai_endpoint = os.environ.get("AZURE_OPENAI_ENDPOINT")
        self.openai_deployment = os.environ.get("AZURE_OPENAI_DEPLOYMENT", "gpt-4o")
        self.openai_key = os.environ.get("AZURE_OPENAI_KEY")
        self.openai_version = os.environ.get("AZURE_OPENAI_VERSION", "2024-11-20")
        
        if self.openai_endpoint:
            try:

                
                self.client = AzureOpenAI(
                    azure_endpoint=self.openai_endpoint,
                    api_key=self.openai_key,
                    api_version=self.openai_version,
                )
                
                # Test the connection
                self._test_connection()
                logging.info("Azure OpenAI client initialized successfully")
                
            except Exception as e:
                logging.error(f"Failed to initialize Azure OpenAI client: {e}")
                self.client = None
        else:
            self.client = None
            logging.warning("Azure OpenAI endpoint not configured. AI-powered parsing will be disabled.")
    
    def _test_connection(self):
        """Test Azure OpenAI connection with a simple request"""
        try:
            # Simple test to verify connectivity
            response = self.client.chat.completions.create(
                model=self.openai_deployment,
                messages=[{"role": "user", "content": "test"}],
                max_tokens=5,
                temperature=0
            )
            logging.info("Azure OpenAI connection test successful")
        except Exception as e:
            logging.warning(f"Azure OpenAI connection test failed: {e}")
            raise
    
    def analyze_order_completeness(self, parsed_data: Dict[str, Any], file_type: str) -> Dict[str, Any]:
        """Analyze order data for completeness using Azure OpenAI"""
        if not self.client:
            return {
                "completeness_score": 0.5,
                "missing_fields": ["AI analysis unavailable"],
                "validation_errors": [],
                "recommendations": ["Configure Azure OpenAI for enhanced validation"]
            }
        
        try:
            # Create a prompt for order analysis
            prompt = self._create_analysis_prompt(parsed_data, file_type)
            
            response = self.client.chat.completions.create(
                model=self.openai_deployment,
                messages=[
                    {"role": "system", "content": "You are an expert in analyzing order data for FMCG supply chain operations. Provide detailed analysis in JSON format."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.1,
                max_tokens=1500,
                response_format={"type": "json_object"}  # Ensure JSON response
            )
            
            analysis_result = json.loads(response.choices[0].message.content)
            return analysis_result
            
        except openai.APIConnectionError as e:
            logging.error(f"Azure OpenAI connection error: {e}")
            return self._get_fallback_analysis("Connection error")
        except openai.RateLimitError as e:
            logging.error(f"Azure OpenAI rate limit error: {e}")
            time.sleep(2)  # Brief backoff
            return self._get_fallback_analysis("Rate limit exceeded")
        except openai.APIStatusError as e:
            logging.error(f"Azure OpenAI API error: {e.status_code} - {e.response}")
            return self._get_fallback_analysis(f"API error: {e.status_code}")
        except json.JSONDecodeError as e:
            logging.error(f"Failed to parse AI response: {e}")
            return self._get_fallback_analysis("Invalid response format")
        except Exception as e:
            logging.error(f"Azure OpenAI analysis failed: {e}")
            return self._get_fallback_analysis(f"Analysis error: {str(e)}")
    
    def extract_sku_items(self, parsed_data: Dict[str, Any], file_type: str) -> List[Dict[str, Any]]:
        """Extract SKU items from parsed data using AI"""
        if not self.client:
            return []
        
        try:
            prompt = self._create_sku_extraction_prompt(parsed_data, file_type)
            
            response = self.client.chat.completions.create(
                model=self.openai_deployment,
                messages=[
                    {"role": "system", "content": "You are an expert in extracting product/SKU information from order data. Return a JSON array of SKU items."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.1,
                max_tokens=2000,
                response_format={"type": "json_object"}  # Ensure JSON response
            )
            
            # Parse response and extract array from JSON object
            response_data = json.loads(response.choices[0].message.content)
            sku_items = response_data.get("sku_items", []) if isinstance(response_data, dict) else response_data
            return sku_items if isinstance(sku_items, list) else []
            
        except openai.APIConnectionError as e:
            logging.error(f"Azure OpenAI connection error during SKU extraction: {e}")
            return []
        except openai.RateLimitError as e:
            logging.error(f"Azure OpenAI rate limit error during SKU extraction: {e}")
            time.sleep(2)  # Brief backoff
            return []
        except openai.APIStatusError as e:
            logging.error(f"Azure OpenAI API error during SKU extraction: {e.status_code} - {e.response}")
            return []
        except json.JSONDecodeError as e:
            logging.error(f"Failed to parse SKU extraction response: {e}")
            return []
        except Exception as e:
            logging.error(f"SKU extraction failed: {e}")
            return []
    
    def _get_fallback_analysis(self, error_reason: str) -> Dict[str, Any]:
        """Return fallback analysis when AI analysis fails"""
        return {
            "completeness_score": 0.5,
            "missing_fields": [f"AI analysis unavailable: {error_reason}"],
            "validation_errors": [],
            "recommendations": ["Manual review required due to AI analysis failure"],
            "order_summary": {
                "total_sku_count": 0,
                "estimated_total_quantity": 0,
                "has_pricing": False,
                "has_delivery_info": False
            }
        }
    
    def _create_analysis_prompt(self, parsed_data: Dict[str, Any], file_type: str) -> str:
        """Create prompt for order completeness analysis"""
        return f"""
        Analyze the following {file_type} order data for completeness and quality:

        Data: {json.dumps(parsed_data, indent=2)}

        Required fields for a complete order:
        - Order identification (order number, reference)
        - Customer/retailer information
        - Delivery details (address, requested delivery date)
        - Product items with SKU codes, quantities, and descriptions
        - Pricing information (unit prices, totals)
        - Special instructions or requirements

        Please provide analysis in this JSON format:
        {{
            "completeness_score": <float between 0.0 and 1.0>,
            "missing_fields": ["list of missing required fields"],
            "validation_errors": ["list of data quality issues"],
            "recommendations": ["list of improvement suggestions"],
            "order_summary": {{
                "total_sku_count": <number>,
                "estimated_total_quantity": <number>,
                "has_pricing": <boolean>,
                "has_delivery_info": <boolean>
            }}
        }}
        """
    
    def _create_sku_extraction_prompt(self, parsed_data: Dict[str, Any], file_type: str) -> str:
        """Create prompt for SKU item extraction"""
        return f"""
        Extract individual SKU/product items from this {file_type} order data:

        Data: {json.dumps(parsed_data, indent=2)}

        Please extract and return a JSON object with a "sku_items" array where each item has this structure:
        {{
            "sku_items": [
                {{
                    "sku_code": "product SKU or identifier",
                    "product_name": "product description/name",
                    "category": "product category if available",
                    "brand": "brand name if available",
                    "quantity_ordered": <integer>,
                    "unit_of_measure": "unit type (pieces, kg, liters, etc.)",
                    "unit_price": <decimal or null>,
                    "total_price": <decimal or null>,
                    "weight_kg": <decimal or null>,
                    "volume_m3": <decimal or null>,
                    "temperature_requirement": "ambient/chilled/frozen or null",
                    "fragile": <boolean>,
                    "product_attributes": {{"any": "additional attributes"}}
                }}
            ]
        }}

        If pricing, weight, or other details are not available, use null values.
        Ensure quantity_ordered is always a positive integer.
        Always return a valid JSON object with the "sku_items" array.
        """

def get_database_connection():
    """Get database connection with proper error handling"""
    try:
        conn = psycopg2.connect(
            host=DB_HOST,
            dbname=DB_NAME,
            user=DB_USER,
            password=DB_PASSWORD,
            port=DB_PORT,
            sslmode='require'
        )
        return conn
    except Exception as e:
        logging.error(f"Database connection failed: {e}")
        raise

def insert_order_tracking(conn, order_id: str, status: str, message: str, details: Dict[str, Any]):
    """Insert order tracking entry"""
    try:
        cur = conn.cursor()
        insert_query = sql.SQL("""
            INSERT INTO order_tracking (id, order_id, status, message, details, created_at)
            VALUES (%s, %s, %s, %s, %s, NOW())
        """)
        cur.execute(insert_query, (
            str(uuid.uuid4()),
            order_id,
            status,
            message,
            json.dumps(details)
        ))
        cur.close()
    except Exception as e:
        logging.error(f"Failed to insert order tracking: {e}")
        raise

def insert_sku_items(conn, order_id: str, sku_items: List[Dict[str, Any]]):
    """Insert SKU items for an order"""
    if not sku_items:
        return
    
    try:
        cur = conn.cursor()
        
        # First, delete existing SKU items for this order
        delete_query = sql.SQL("DELETE FROM order_sku_items WHERE order_id = %s")
        cur.execute(delete_query, (order_id,))
        
        # Insert new SKU items
        insert_query = sql.SQL("""
            INSERT INTO order_sku_items (
                id, order_id, sku_code, product_name, category, brand,
                quantity_ordered, unit_of_measure, unit_price, total_price,
                weight_kg, volume_m3, temperature_requirement, fragile,
                product_attributes, created_at, updated_at
            ) VALUES (
                %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, NOW(), NOW()
            )
        """)
        
        for item in sku_items:
            cur.execute(insert_query, (
                str(uuid.uuid4()),
                order_id,
                item.get('sku_code', ''),
                item.get('product_name', ''),
                item.get('category'),
                item.get('brand',
                item.get('quantity_ordered', 0),
                item.get('unit_of_measure'),
                item.get('unit_price'),
                item.get('total_price'),
                item.get('weight_kg'),
                item.get('volume_m3'),
                item.get('temperature_requirement'),
                item.get('fragile', False),
                json.dumps(item.get('product_attributes', {})),
            )))
        
        cur.close()
        logging.info(f"Inserted {len(sku_items)} SKU items for order {order_id}")
        
    except Exception as e:
        logging.error(f"Failed to insert SKU items: {e}")
        raise

def update_order_summary(conn, order_id: str, analysis_result: Dict[str, Any], sku_items: List[Dict[str, Any]]):
    """Update order with analysis results and summary"""
    try:
        cur = conn.cursor()
        
        # Calculate summary statistics
        total_sku_count = len(sku_items)
        total_quantity = sum(item.get('quantity_ordered', 0) for item in sku_items)
        total_weight = sum(item.get('weight_kg', 0) or 0 for item in sku_items)
        total_volume = sum(item.get('volume_m3', 0) or 0 for item in sku_items)
        subtotal = sum(item.get('total_price', 0) or 0 for item in sku_items)
        
        update_query = sql.SQL("""
            UPDATE orders SET
                parsed_data = %s,
                missing_fields = %s,
                validation_errors = %s,
                total_sku_count = %s,
                total_quantity = %s,
                total_weight_kg = %s,
                total_volume_m3 = %s,
                subtotal = %s,
                updated_at = NOW()
            WHERE id = %s
        """)
        
        cur.execute(update_query, (
            json.dumps(analysis_result.get('order_summary', {})),
            json.dumps(analysis_result.get('missing_fields', [])),
            json.dumps(analysis_result.get('validation_errors', [])),
            total_sku_count,
            total_quantity,
            total_weight,
            total_volume,
            subtotal,
            order_id
        ))
        
        cur.close()
        
    except Exception as e:
        logging.error(f"Failed to update order summary: {e}")
        raise

@app.route(route="order_file_reader")
def order_file_reader(req: func.HttpRequest) -> func.HttpResponse:
    """Enhanced order file reader with AI-powered parsing and completeness checking"""
    logging.info('Enhanced order file reader function processed a request.')
    
    # Get order_id from request
    order_id = req.params.get('order_id')
    if not order_id:
        try:
            req_body = req.get_json()
        except ValueError:
            req_body = None
        if req_body:
            order_id = req_body.get('order_id')

    if not order_id:
        return func.HttpResponse(
            "Please provide order_id in the query string or in the request body.",
            status_code=400
        )

    # Initialize services
    parsing_service = OrderParsingService()
    
    # Query order details from database
    order_details = None
    parse_status = "PARSING_COMPLETE"
    parse_message = ""
    parsed_data = {}
    analysis_result = {}
    sku_items = []
    
    try:
        conn = get_database_connection()
        cur = conn.cursor()
        
        # Query order details using new schema
        query = sql.SQL("""
            SELECT id, order_number, priority, requested_delivery_date, file_path, 
                   status, original_filename, file_type, file_size
            FROM orders
            WHERE id = %s
        """)
        cur.execute(query, (order_id,))
        row = cur.fetchone()
        order_details = row

        if row and row[4]:  # file_path exists
            order_uuid, order_number, priority, delivery_date, file_path, status, filename, file_type, file_size = row
            blob_connection_str = os.environ.get("AzureWebJobsStorage")
            
            try:
                # Validate file_path format
                if "/" not in file_path:
                    raise ValueError("file_path must be in the format 'container/blobname'")

                container = "requestedorders"
                _, blob_name = file_path.split(f"{container}/", 1)
                blob_service_client = BlobServiceClient.from_connection_string(blob_connection_str)
                blob_client = blob_service_client.get_blob_client(container=container, blob=blob_name)
                logging.info(f"Processing file: {blob_name} in container: {container}")
                
                # Get file information
                file_extension = os.path.splitext(blob_name)[1].lower()
                blob_props = blob_client.get_blob_properties()
                content_type = blob_props.content_settings.content_type
                
                # Download and parse file content
                file_data = blob_client.download_blob().readall()
                parsed_data = parse_file_content(file_data, file_extension, blob_name)
                
                if parsed_data.get("file_type") != "Unsupported":
                    # Use AI to analyze order completeness
                    analysis_result = parsing_service.analyze_order_completeness(
                        parsed_data, parsed_data.get("file_type", "Unknown")
                    )
                    
                    # Extract SKU items using AI
                    sku_items = parsing_service.extract_sku_items(
                        parsed_data, parsed_data.get("file_type", "Unknown")
                    )
                    
                    # Update parse status based on analysis
                    completeness_score = analysis_result.get('completeness_score', 0.0)
                    if completeness_score >= 0.8:
                        parse_status = "PARSING_COMPLETE"
                        parse_message = f"Order parsed successfully with {completeness_score:.1%} completeness"
                    elif completeness_score >= 0.5:
                        parse_status = "PARSING_PARTIAL"
                        parse_message = f"Order parsed with {completeness_score:.1%} completeness. Some fields may be missing."
                    else:
                        parse_status = "PARSING_INCOMPLETE"
                        parse_message = f"Order parsing incomplete ({completeness_score:.1%}). Manual review required."
                    
                    # Insert SKU items into database
                    insert_sku_items(conn, order_id, sku_items)
                    
                    # Update order with analysis results
                    update_order_summary(conn, order_id, analysis_result, sku_items)
                    
                else:
                    parse_status = "PARSING_FAILED"
                    parse_message = f"Unsupported file type: {file_extension}"
                    
            except Exception as e:
                parse_status = "PARSING_FAILED"
                parse_message = f"Failed to parse file: {e}"
                logging.error(parse_message)
                parsed_data = {"error": str(e)}
                
        else:
            parse_status = "PARSING_FAILED"
            parse_message = "No file_path found for this order or order not found."

        # Insert parsing results into order_tracking
        details = {
            "parsing_results": parsed_data,
            "ai_analysis": analysis_result,
            "sku_items_count": len(sku_items),
            "file_path": file_path if 'file_path' in locals() else None
        }
        
        insert_order_tracking(conn, order_id, parse_status, parse_message, details)
        conn.commit()
        cur.close()
        conn.close()
        
    except Exception as e:
        logging.error(f"Database error: {e}")
        return func.HttpResponse(
            f"Database error: {e}",
            status_code=500
        )

    # Prepare response
    if order_details:
        order_uuid, order_number, priority, delivery_date, file_path, status, filename, file_type, file_size = order_details
        
        result = {
            "order_id": order_id,
            "order_number": order_number,
            "priority": priority,
            "requested_delivery_date": delivery_date.strftime('%Y-%m-%d') if isinstance(delivery_date, datetime) else str(delivery_date) if delivery_date else None,
            "file_info": {
                "original_filename": filename,
                "file_path": file_path,
                "file_type": file_type,
                "file_size": file_size
            },
            "parse_status": parse_status,
            "parse_message": parse_message,
            "parsed_data": parsed_data,
            "ai_analysis": analysis_result,
            "sku_items": sku_items,
            "summary": {
                "total_sku_count": len(sku_items),
                "completeness_score": analysis_result.get('completeness_score', 0.0),
                "missing_fields_count": len(analysis_result.get('missing_fields', [])),
                "validation_errors_count": len(analysis_result.get('validation_errors', []))
            }
        }
        
        return func.HttpResponse(
            json.dumps(result, default=str),
            status_code=200,
            mimetype="application/json"
        )
    else:
        return func.HttpResponse(
            f"Order with id {order_id} not found.",
            status_code=404
        )

def parse_file_content(file_data: bytes, file_extension: str, blob_name: str) -> Dict[str, Any]:
    """Parse file content based on file type"""
    try:
        if file_extension in ['.csv']:
            # Parse CSV
            import io
            df = pd.read_csv(io.BytesIO(file_data))
            return {
                "file_type": "CSV",
                "columns": df.columns.tolist(),
                "row_count": len(df),
                "sample_data": df.head(10).to_dict('records'),
                "data_types": df.dtypes.astype(str).to_dict(),
                "all_data": df.to_dict('records')  # Include all data for AI analysis
            }
            
        elif file_extension in ['.xlsx', '.xls']:
            # Parse Excel
            import io
            df = pd.read_excel(io.BytesIO(file_data))
            return {
                "file_type": "Excel",
                "columns": df.columns.tolist(),
                "row_count": len(df),
                "sample_data": df.head(10).to_dict('records'),
                "data_types": df.dtypes.astype(str).to_dict(),
                "all_data": df.to_dict('records')  # Include all data for AI analysis
            }
            
        elif file_extension in ['.json']:
            # Parse JSON
            json_data = json.loads(file_data.decode('utf-8'))
            if isinstance(json_data, list):
                return {
                    "file_type": "JSON",
                    "structure": "array",
                    "item_count": len(json_data),
                    "sample_data": json_data[:10] if len(json_data) > 0 else [],
                    "keys": list(json_data[0].keys()) if len(json_data) > 0 and isinstance(json_data[0], dict) else [],
                    "all_data": json_data
                }
            else:
                return {
                    "file_type": "JSON",
                    "structure": "object",
                    "keys": list(json_data.keys()) if isinstance(json_data, dict) else [],
                    "sample_data": json_data,
                    "all_data": json_data
                }
                
        elif file_extension in ['.txt']:
            # Parse text file
            text_content = file_data.decode('utf-8', errors='replace')
            lines = text_content.split('\n')
            return {
                "file_type": "Text",
                "line_count": len(lines),
                "character_count": len(text_content),
                "word_count": len(text_content.split()),
                "sample_lines": lines[:20],
                "full_content": text_content
            }
            
        elif file_extension in ['.docx']:
            # Parse Word document
            import io
            doc = Document(io.BytesIO(file_data))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            tables_data = []
            
            # Extract table data
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables_data.append(table_data)
            
            return {
                "file_type": "Word Document",
                "paragraph_count": len(paragraphs),
                "sample_paragraphs": paragraphs[:10],
                "tables_count": len(doc.tables),
                "tables_data": tables_data,
                "full_text": '\n'.join(paragraphs)
            }
            
        else:
            # Unsupported file type
            return {
                "file_type": "Unsupported",
                "file_extension": file_extension,
                "file_size": len(file_data)
            }
            
    except Exception as e:
        logging.error(f"Error parsing file content: {e}")
        return {
            "file_type": "Error",
            "error": str(e),
            "file_extension": file_extension
        }

@app.route(route="validate_order_completeness")
def validate_order_completeness(req: func.HttpRequest) -> func.HttpResponse:
    """Validate order completeness and provide recommendations"""
    logging.info('Order completeness validation function processed a request.')
    
    # Get order_id from request
    order_id = req.params.get('order_id')
    if not order_id:
        try:
            req_body = req.get_json()
        except ValueError:
            req_body = None
        if req_body:
            order_id = req_body.get('order_id')

    if not order_id:
        return func.HttpResponse(
            "Please provide order_id in the query string or in the request body.",
            status_code=400
        )

    try:
        conn = get_database_connection()
        cur = conn.cursor()
        
        # Get order details with parsed data
        query = sql.SQL("""
            SELECT o.id, o.order_number, o.parsed_data, o.missing_fields, 
                   o.validation_errors, o.total_sku_count, o.status,
                   COUNT(osi.id) as actual_sku_count
            FROM orders o
            LEFT JOIN order_sku_items osi ON o.id = osi.order_id
            WHERE o.id = %s
            GROUP BY o.id, o.order_number, o.parsed_data, o.missing_fields, 
                     o.validation_errors, o.total_sku_count, o.status
        """)
        cur.execute(query, (order_id,))
        row = cur.fetchone()
        
        if not row:
            cur.close()
            conn.close()
            return func.HttpResponse(
                f"Order with id {order_id} not found.",
                status_code=404
            )
        
        order_uuid, order_number, parsed_data, missing_fields, validation_errors, total_sku_count, status, actual_sku_count = row
        
        # Get SKU items
        sku_query = sql.SQL("""
            SELECT sku_code, product_name, quantity_ordered, unit_price, 
                   category, brand, product_attributes
            FROM order_sku_items
            WHERE order_id = %s
        """)
        cur.execute(sku_query, (order_id,))
        sku_rows = cur.fetchall()
        
        cur.close()
        conn.close()
        
        # Initialize parsing service for AI validation
        parsing_service = OrderParsingService()
        
        # Prepare data for AI analysis
        current_data = {
            "order_info": {
                "order_number": order_number,
                "status": status,
                "total_sku_count": total_sku_count,
                "actual_sku_count": actual_sku_count
            },
            "parsed_data": json.loads(parsed_data) if parsed_data else {},
            "sku_items": [
                {
                    "sku_code": row[0],
                    "product_name": row[1],
                    "quantity_ordered": row[2],
                    "unit_price": row[3],
                    "category": row[4],
                    "brand": row[5],
                    "product_attributes": json.loads(row[6]) if row[6] else {}
                }
                for row in sku_rows
            ]
        }
        
        # Perform AI-powered completeness analysis
        analysis_result = parsing_service.analyze_order_completeness(
            current_data, "Order Validation"
        )
        
        # Calculate completeness metrics
        completeness_metrics = {
            "has_order_number": bool(order_number),
            "has_sku_items": actual_sku_count > 0,
            "sku_count_match": total_sku_count == actual_sku_count,
            "has_pricing": any(item.get("unit_price") for item in current_data["sku_items"]),
            "has_categories": any(item.get("category") for item in current_data["sku_items"]),
            "completeness_score": analysis_result.get("completeness_score", 0.0)
        }
        
        # Generate recommendations
        recommendations = analysis_result.get("recommendations", [])
        if not completeness_metrics["has_sku_items"]:
            recommendations.append("No SKU items found. Please ensure the order file contains product information.")
        if not completeness_metrics["sku_count_match"]:
            recommendations.append(f"SKU count mismatch: Expected {total_sku_count}, found {actual_sku_count}")
        if not completeness_metrics["has_pricing"]:
            recommendations.append("Missing pricing information for SKU items.")
        
        # Prepare response
        result = {
            "order_id": order_id,
            "order_number": order_number,
            "validation_timestamp": datetime.now().isoformat(),
            "completeness_metrics": completeness_metrics,
            "ai_analysis": analysis_result,
            "current_missing_fields": json.loads(missing_fields) if missing_fields else [],
            "current_validation_errors": json.loads(validation_errors) if validation_errors else [],
            "recommendations": recommendations,
            "sku_summary": {
                "total_items": actual_sku_count,
                "items_with_pricing": sum(1 for item in current_data["sku_items"] if item.get("unit_price")),
                "items_with_categories": sum(1 for item in current_data["sku_items"] if item.get("category")),
                "unique_brands": len(set(item.get("brand") for item in current_data["sku_items"] if item.get("brand")))
            },
            "next_steps": get_next_steps(completeness_metrics, analysis_result)
        }
        
        return func.HttpResponse(
            json.dumps(result, default=str),
            status_code=200,
            mimetype="application/json"
        )
        
    except Exception as e:
        logging.error(f"Order validation error: {e}")
        return func.HttpResponse(
            f"Order validation error: {e}",
            status_code=500
        )

def get_next_steps(metrics: Dict[str, Any], analysis: Dict[str, Any]) -> List[str]:
    """Generate next steps based on completeness analysis"""
    steps = []
    
    completeness_score = metrics.get("completeness_score", 0.0)
    
    if completeness_score >= 0.9:
        steps.append("Order is highly complete and ready for processing")
        steps.append("Consider moving to 'READY_FOR_ASSIGNMENT' status")
    elif completeness_score >= 0.7:
        steps.append("Order is mostly complete with minor issues")
        steps.append("Review missing fields and validation errors")
        steps.append("Consider manual review before processing")
    elif completeness_score >= 0.5:
        steps.append("Order has significant gaps that need attention")
        steps.append("Contact customer for missing information")
        steps.append("Hold order in 'PENDING_REVIEW' status")
    else:
        steps.append("Order is incomplete and requires major revision")
        steps.append("Request customer to resubmit with complete information")
        steps.append("Set status to 'INCOMPLETE' and notify customer")
    
    # Add specific steps based on missing elements
    if not metrics.get("has_sku_items"):
        steps.append("Critical: Add product/SKU information")
    if not metrics.get("has_pricing"):
        steps.append("Important: Add pricing information for cost calculations")
    
    return steps

@app.route(route="health")
def health_check(req: func.HttpRequest) -> func.HttpResponse:
    """Health check endpoint"""
    return func.HttpResponse(
        json.dumps({
            "status": "healthy",
            "timestamp": datetime.now().isoformat(),
            "azure_openai_configured": bool(os.environ.get("AZURE_OPENAI_ENDPOINT"))
        }),
        status_code=200,
        mimetype="application/json"
    )